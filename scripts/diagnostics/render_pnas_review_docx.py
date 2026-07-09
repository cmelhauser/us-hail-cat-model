#!/usr/bin/env python3
"""
Build a double-spaced, line-numbered Word review draft with embedded figures.

Reads ``docs/pnas_article_publication.md`` (regenerate that file first if stale)
and writes ``docs/pnas_article_review.docx``.

Usage (repo root):
  .venv/bin/python scripts/diagnostics/render_pnas_review_docx.py
"""

from __future__ import annotations

import subprocess
import sys
import tempfile
from datetime import date
from pathlib import Path

REPO = Path(__file__).resolve().parents[2]
if str(REPO) not in sys.path:
    sys.path.insert(0, str(REPO))

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from scripts.diagnostics.render_pnas_publication_md import MANUSCRIPT_TITLE

DOCS = REPO / "docs"
SRC_MD = DOCS / "pnas_article_publication.md"
OUT_DOCX = DOCS / "pnas_article_review.docx"
FIG_DIR = DOCS / "figures" / "pnas"


def _enable_line_numbers(doc: Document) -> None:
    for section in doc.sections:
        sect_pr = section._sectPr
        for child in list(sect_pr):
            if child.tag == qn("w:lnNumType"):
                sect_pr.remove(child)
        ln_num = OxmlElement("w:lnNumType")
        ln_num.set(qn("w:countBy"), "1")
        ln_num.set(qn("w:restart"), "continuous")
        ln_num.set(qn("w:distance"), "360")
        sect_pr.append(ln_num)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.0)


def _apply_double_spacing(doc: Document) -> None:
    for style_name in ("Normal", "Body Text", "First Paragraph", "Compact"):
        try:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.space_before = Pt(0)
        except KeyError:
            pass

    for para in doc.paragraphs:
        pf = para.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        for run in para.runs:
            if not run.font.name:
                run.font.name = "Times New Roman"
            if run.font.size is None:
                run.font.size = Pt(12)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE


def _pending_figure_note() -> str:
    expected = [
        "fig12_rp_100yr_stochastic.png",
        "fig13_analytical_vs_stochastic.png",
    ]
    missing = [f for f in expected if not (FIG_DIR / f).is_file()]
    if not missing:
        return ""
    names = ", ".join(missing)
    return (
        f"\n\n> **Note:** Figures pending Stage 13/14 outputs: {names}. "
        "Re-run `render_pnas_article_figures.py` after the stochastic catalog completes.\n"
    )


def _prepare_markdown() -> str:
    if not SRC_MD.is_file():
        raise FileNotFoundError(
            f"Missing {SRC_MD}; run render_pnas_publication_md.py first."
        )
    body = SRC_MD.read_text()
    # Pandoc resolves image paths relative to the markdown file location (docs/).
    header = f"""% {MANUSCRIPT_TITLE}

**DRAFT FOR REVIEW — {date.today().isoformat()}**  
Double-spaced manuscript with continuous line numbers and embedded figures.  
Model repository: [github.com/cmelhauser/us-hail-cat-model](https://github.com/cmelhauser/us-hail-cat-model)

---

"""
    note = _pending_figure_note()
    if note and "## Figures" in body:
        body = body.replace("## Figures", note + "\n## Figures", 1)
    return header + body


def _pandoc_to_docx(md_text: str, out_path: Path) -> None:
    with tempfile.NamedTemporaryFile(
        mode="w",
        suffix=".md",
        delete=False,
        encoding="utf-8",
        dir=DOCS,
    ) as tmp:
        tmp.write(md_text)
        tmp_path = Path(tmp.name)

    try:
        cmd = [
            "pandoc",
            str(tmp_path),
            "-f",
            "markdown",
            "-t",
            "docx",
            "-o",
            str(out_path),
            "--standalone",
        ]
        subprocess.run(cmd, check=True, cwd=DOCS)
    finally:
        tmp_path.unlink(missing_ok=True)


def _postprocess_docx(path: Path) -> None:
    doc = Document(path)
    _enable_line_numbers(doc)
    _apply_double_spacing(doc)
    doc.save(path)


def main() -> None:
    md_text = _prepare_markdown()
    _pandoc_to_docx(md_text, OUT_DOCX)
    _postprocess_docx(OUT_DOCX)
    n_figs = len(list(FIG_DIR.glob("fig*.png")))
    print(f"Wrote {OUT_DOCX} ({n_figs} figure PNGs on disk; embedded where referenced in markdown)")


if __name__ == "__main__":
    main()
